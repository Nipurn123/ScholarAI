from docx import Document
from docx.shared import Inches

def generate_word_report(papers, filename):
    doc = Document()
    doc.add_heading('Publication Report', 0)
    
    for paper in papers:
        doc.add_paragraph(f"Title: {paper['title']}")
        doc.add_paragraph(f"Authors: {', '.join(paper['authors'])}")
        doc.add_paragraph(f"Year: {paper['year']}")
        doc.add_paragraph(f"Venue: {paper['venue']}")
        doc.add_paragraph(f"Citations: {paper['citations']}")
        doc.add_paragraph(f"Type: {paper['type']}")
        doc.add_paragraph('---')
    
    doc.save(filename)
    print(f"Word report generated: {filename}")