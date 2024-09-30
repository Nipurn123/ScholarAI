import streamlit as st
from searxng import search_searxng
from datetime import datetime
from data_acquisition_module.google_scholar_crawler import scrape_google_scholar_profile
from data_processing_module.organiser import organize_papers_by_year_and_citations
from data_processing_module.duplicate_remover import remove_duplicates
from data_processing_module.data_normalizer import normalize_data
from data_processing_module.date_range_filter import filter_by_date
from langchain_groq import ChatGroq
import pandas as pd
from tenacity import retry, stop_after_attempt, wait_exponential
import io
import bibtexparser
import openpyxl
import os
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Font
from openpyxl.utils import column_index_from_string
from docx import Document
from docx.shared import Inches

@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
def categorize_paper(llm, title, venue):
    messages = [
        (
            "system",
            "You are an advanced assistant designed to categorize research papers based on their titles and publication venues into two distinct domains: 'Journal Research Papers' and 'Conference Research Papers.'"
        ),
        (
            "human",
            f"Categorize this paper: Title: {title}, Venue: {venue}. Respond with only 'Journal' or 'Conference'."
        ),
    ]
    
    ai_msg = llm.invoke(messages)
    return ai_msg.content.strip().lower()

def categorize_papers(papers, llm):
    journal_papers = []
    conference_papers = []

    for _, paper in papers.iterrows():
        try:
            category = categorize_paper(llm, paper['title'], paper['venue'])
            if category == 'journal':
                journal_papers.append(paper.to_dict())
            elif category == 'conference':
                conference_papers.append(paper.to_dict())
        except Exception as e:
            st.warning(f"Error categorizing paper: {paper['title']}. Error: {str(e)}")

    return journal_papers, conference_papers

def save_to_word(data, filename, title):
    doc = Document()
    doc.add_heading(title, 0)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Year'
    hdr_cells[1].text = 'Citations'
    hdr_cells[2].text = 'Title'
    hdr_cells[3].text = 'Authors'
    hdr_cells[4].text = 'Venue'

    for _, row in data.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['year'])
        row_cells[1].text = str(row['citations'])
        row_cells[2].text = row['title']
        row_cells[3].text = row['authors']
        row_cells[4].text = row['venue']

    doc.save(filename)

def save_categorized_papers(journal_papers, conference_papers, faculty_name, start_year, end_year):
    journal_df = pd.DataFrame(journal_papers)
    conference_df = pd.DataFrame(conference_papers)

    # Sort dataframes by year (descending) and citations
    journal_df = journal_df.sort_values(by=['year', 'citations'], ascending=[False, False])
    conference_df = conference_df.sort_values(by=['year', 'citations'], ascending=[False, False])

    # Save to Excel
    excel_filename = f"{faculty_name}_papers_{start_year}_{end_year}.xlsx"
    with pd.ExcelWriter(excel_filename, engine='openpyxl') as writer:
        journal_df.to_excel(writer, sheet_name='Journal Papers', index=False)
        conference_df.to_excel(writer, sheet_name='Conference Papers', index=False)

    # Save to Word
    journal_word_filename = f"{faculty_name}_journal_papers_{start_year}_{end_year}.docx"
    conference_word_filename = f"{faculty_name}_conference_papers_{start_year}_{end_year}.docx"
    
    save_to_word(journal_df, journal_word_filename, f"{faculty_name} Journal Papers ({start_year}-{end_year})")
    save_to_word(conference_df, conference_word_filename, f"{faculty_name} Conference Papers ({start_year}-{end_year})")

    return excel_filename, journal_word_filename, conference_word_filename

def extract_names_from_bibtex(bibtex_string):
    bib_database = bibtexparser.loads(bibtex_string)
    names = set()
    for entry in bib_database.entries:
        if 'author' in entry:
            authors = entry['author'].split(' and ')
            for author in authors:
                names.add(author.strip())
    return list(names)

def read_names_from_excel(uploaded_file):
    df = pd.read_excel(uploaded_file)
    if 'Faculty Name' in df.columns:
        return df['Faculty Name'].dropna().tolist()
    else:
        st.error("The Excel file must contain a column named 'Faculty Name'")
        return []

def save_raw_data_to_excel(papers, faculty_name):
    df = pd.DataFrame(papers)
    filename = f"{faculty_name}_raw_data.xlsx"
    df.to_excel(filename, index=False)
    return filename

def process_raw_data(filename, start_year, end_year):
    df = pd.read_excel(filename)
    
    # Remove duplicates
    df = df.drop_duplicates(subset=['title'], keep='first')
    
    # Normalize data (assuming this involves some column operations)
    df['year'] = pd.to_datetime(df['year'], format='%Y').dt.year
    
    # Filter by date range
    df = df[(df['year'] >= start_year) & (df['year'] <= end_year)]
    
    # Sort by year (descending) and citations
    df = df.sort_values(by=['year', 'citations'], ascending=[False, False])
    
    return df

def process_custom_query(llm, query, start_year, end_year):
    messages = [
        (
            "system",
            "You are an assistant designed to create Excel formulas for custom queries on research paper data. The data includes columns for year (A), citations (B), title (C), authors (D), and venue (E). Generate an Excel formula to extract the relevant data based on the user's query and specified year range. Provide only the Excel formula, nothing else."
        ),
        (
            "human",
            f"Create an Excel formula for the following query: '{query}' and year range {start_year}-{end_year}. The data is in columns A (year), B (citations), C (title), D (authors), E (venue). Return only the formula, no explanations or notes."
        ),
    ]
    
    ai_msg = llm.invoke(messages)
    return ai_msg.content.strip()

def apply_excel_formula(excel_file, sheet_name, formula):
    wb = openpyxl.load_workbook(excel_file)
    sheet = wb[sheet_name]
    
    # Find the last row with data
    last_row = sheet.max_row
    
    # Apply the formula to the range
    for row in range(2, last_row + 1):
        cell = sheet.cell(row=row, column=6)  # Column F
        cell.value = f"={formula.replace('2', str(row))}"
    
    # Save the workbook
    wb.save(excel_file)
    return last_row

def extract_formula_results(sheet):
    results = []
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if row[5]:  # If the formula result in column F is True
            results.append(row[:5])  # Append the first 5 columns (A-E)
    
    # Create a DataFrame with the extracted data
    columns = ['Year', 'Citations', 'Title', 'Authors', 'Venue']
    df = pd.DataFrame(results, columns=columns)
    
    return df

def main():
    st.set_page_config(page_title="Faculty Research Analysis", layout="wide")
    
    st.title("Faculty Research Analysis")
    st.markdown("""
    This app analyzes research papers for selected faculty members.
    You can customize the year range and provide a custom query after data retrieval.
    """)
    
    # Input type selection
    input_type = st.radio("Select input type:", ("BibTeX", "Excel"))
    
    selected_names = []
    
    if input_type == "BibTeX":
        bibtex_input = st.text_area("Enter BibTeX string:")
        if bibtex_input:
            try:
                names = extract_names_from_bibtex(bibtex_input)
                selected_names = st.multiselect("Select faculty names:", names)
            except Exception as e:
                st.error(f"Error processing BibTeX: {str(e)}")
    else:  # Excel
        uploaded_file = st.file_uploader("Upload Excel file", type=["xlsx", "xls"])
        if uploaded_file is not None:
            names = read_names_from_excel(uploaded_file)
            selected_names = st.multiselect("Select faculty names:", names)
    
    if not selected_names:
        st.warning("Please select at least one faculty name to proceed.")
        return
    
    retrieve_data_button = st.button("Retrieve Data for Selected Faculty")
    
    if retrieve_data_button:
        for faculty_name in selected_names:
            st.subheader(f"Retrieving data for: {faculty_name}")
            
            # Search for the faculty's Google Scholar profile
            search_query = f"{faculty_name} Google Scholar"
            with st.spinner(f"Searching for {faculty_name}'s profile..."):
                search_results = search_searxng(search_query)
            
            if search_results['results']:
                # Extract the first result's URL for further analysis
                profile_url = search_results['results'][0].url
                
                st.write(f"Found profile: {profile_url}")
                
                try:
                    # Scrape Google Scholar profile and save raw data
                    with st.spinner("Scraping profile and saving raw data..."):
                        papers = scrape_google_scholar_profile(profile_url)
                        raw_data_file = save_raw_data_to_excel(papers, faculty_name)
                    
                    st.success(f"Data retrieved and stored for {faculty_name}.")
                    
                    # Provide download link for the raw Excel file
                    st.download_button(
                        label=f"Download Raw Data Excel for {faculty_name}",
                        data=open(raw_data_file, 'rb').read(),
                        file_name=f"{faculty_name}_raw_data.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                
                except Exception as e:
                    st.error(f"An error occurred during data retrieval for {faculty_name}: {str(e)}")
            
            else:
                st.warning(f"No profile found for {faculty_name}.")
        
        st.session_state.data_retrieved = True
    
    if 'data_retrieved' in st.session_state and st.session_state.data_retrieved:
        st.subheader("Data Analysis Options")
        
        # Year range selection
        st.subheader("Customize Year Range")
        current_year = datetime.now().year
        start_year = st.number_input("Start Year", min_value=1900, max_value=current_year, value=2005)
        end_year = st.number_input("End Year", min_value=1900, max_value=current_year, value=current_year)
        
        if start_year > end_year:
            st.error("Start year cannot be greater than end year.")
            return
        
        # Custom query input
        st.subheader("Custom Query")
        custom_query = st.text_input("Enter your custom query:")
        
        analyze_button = st.button("Analyze Data")
        
        if analyze_button:
            # Initialize LLM
            llm = ChatGroq(
                model="mixtral-8x7b-32768",
                temperature=0,
                max_tokens=None,
                timeout=None,
                max_retries=2,
            )
            
            for faculty_name in selected_names:
                st.subheader(f"Analyzing data for: {faculty_name}")
                
                try:
                    # Process raw data
                    with st.spinner("Processing raw data..."):
                        processed_df = process_raw_data(f"{faculty_name}_raw_data.xlsx", start_year, end_year)
                    
                    if processed_df.empty:
                        st.warning(f"No papers found in the year range {start_year}-{end_year} for {faculty_name}.")
                    else:
                        # Categorize papers using LLM
                        with st.spinner("Categorizing papers..."):
                            journal_papers, conference_papers = categorize_papers(processed_df, llm)
                        
                        # Save categorized papers to Excel and Word
                        excel_file, journal_word_file, conference_word_file = save_categorized_papers(
                            journal_papers, conference_papers, faculty_name, start_year, end_year
                        )
                        
                        st.success(f"Data processing and categorization complete for {faculty_name}.")
                        
                        # Display summary
                        st.write(f"Total papers: {len(journal_papers) + len(conference_papers)}")
                        st.write(f"Journal papers: {len(journal_papers)}")
                        st.write(f"Conference papers: {len(conference_papers)}")
                        
                        # Apply custom query if provided
                        if custom_query:
                            try:
                                # Generate Excel formula
                                excel_formula = process_custom_query(llm, custom_query, start_year, end_year)
                                st.subheader("Generated Excel Formula")
                                st.code(excel_formula, language="excel")

                                # Apply the formula to both Journal and Conference sheets
                                journal_last_row = apply_excel_formula(excel_file, "Journal Papers", excel_formula)
                                conference_last_row = apply_excel_formula(excel_file, "Conference Papers", excel_formula)

                                # Reload the workbook to get the updated values
                                wb = openpyxl.load_workbook(excel_file, data_only=True)
                                
                                # Extract the results
                                journal_results = extract_formula_results(wb["Journal Papers"])
                                conference_results = extract_formula_results(wb["Conference Papers"])

                                st.success(f"Custom query applied and data extracted for {faculty_name}.")
                                
                                # Display results
                                st.subheader("Journal Papers Results")
                                st.write(journal_results)
                                
                                st.subheader("Conference Papers Results")
                                st.write(conference_results)

                                # Provide download buttons for the extracted data
                                journal_csv = journal_results.to_csv(index=False).encode('utf-8')
                                conference_csv = conference_results.to_csv(index=False).encode('utf-8')

                                st.download_button(
                                    label=f"Download Journal Results CSV for {faculty_name}",
                                    data=journal_csv,
                                    file_name=f"{faculty_name}_journal_results.csv",
                                    mime="text/csv"
                                )

                                st.download_button(
                                    label=f"Download Conference Results CSV for {faculty_name}",
                                    data=conference_csv,
                                    file_name=f"{faculty_name}_conference_results.csv",
                                    mime="text/csv"
                                )
                            
                            except Exception as e:
                                st.error(f"An error occurred while applying the custom query: {str(e)}")
                        
                        # Provide download buttons for Excel and Word files
                        st.download_button(
                            label=f"Download Analyzed Papers Excel for {faculty_name}",
                            data=open(excel_file, 'rb').read(),
                            file_name=excel_file,
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                        
                        st.download_button(
                            label=f"Download Journal Papers Word Document for {faculty_name}",
                            data=open(journal_word_file, 'rb').read(),
                            file_name=journal_word_file,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        st.download_button(
                            label=f"Download Conference Papers Word Document for {faculty_name}",
                            data=open(conference_word_file, 'rb').read(),
                            file_name=conference_word_file,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                
                except Exception as e:
                    st.error(f"An error occurred during the analysis for {faculty_name}: {str(e)}")

    # Footer
    st.markdown("---")
    st.markdown("Powered by SearXNG, Google Scholar, and Groq | Created with Streamlit")

if __name__ == "__main__":
    main()
